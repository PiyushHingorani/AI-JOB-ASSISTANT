import re
import pdfplumber
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, Inches

def extract_raw_text_from_pdf(file_path):
    resume_text = ''
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            resume_text += page.extract_text() + ' '
    return resume_text

def clean_text(text):
    # Remove HTML tags
    text = re.sub(r'<[^>]*?>', '', text)
    # Remove URLs
    text = re.sub(r'http[s]?://(?:[a-zA-Z]|[0-9]|[$-_@.&+]|[!*\\(\\),]|(?:%[0-9a-fA-F][0-9a-fA-F]))+', '', text)
    # Remove special characters
    text = re.sub(r'[^a-zA-Z0-9 ]', '', text)
    # Replace multiple spaces with a single space
    text = re.sub(r'\s{2,}', ' ', text)
    # Trim leading and trailing whitespace
    text = text.strip()
    return text

def chunk_text(text, max_tokens=500):
    """
    Splits text into chunks of approximately max_tokens words.
    Adjust max_tokens based on your LLM's context window.
    """
    words = text.split()
    for i in range(0, len(words), max_tokens):
        yield ' '.join(words[i:i + max_tokens])

def create_formatted_cover_letter_docx(cover_letter_text, filename='cover_letter.docx'):
    # Create a new Document
    doc = Document()
    
    # Adjust left and right margins
    sections = doc.sections
    for section in sections:
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    # Split the cover letter into paragraphs
    cover_letter_sections = cover_letter_text.split("\n\n")
    length = len(cover_letter_sections)

    if length >= 3:
        header = doc.add_paragraph()
        header.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        first_bold = False
        for line in cover_letter_sections[0].split("\n"):
            header.add_run(f"{line}\n").bold = not first_bold
            first_bold = True
        header.style.font.size = Pt(14)
        header.style.font.name = "Calibri"

        recipient = doc.add_paragraph()
        recipient.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        for line in cover_letter_sections[1].split("\n"):
            recipient.add_run(f"{line}\n").bold
            first_bold = True
        recipient.style.font.size = Pt(14)
        
        salutation = doc.add_paragraph(f"{cover_letter_sections[2]}\n")
        salutation.style.font.size = Pt(12)

        for section in cover_letter_sections[3:length-1]:
            paragraph_text = section.replace("\n", ' ')
            paragraph = doc.add_paragraph(paragraph_text)
            paragraph.style.font.size = Pt(12)
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

        greetings = doc.add_paragraph(f"{cover_letter_sections[length-1]}\n")
        greetings.style.font.size = Pt(12)
    else:
        paragraph = doc.add_paragraph(cover_letter_text)
        paragraph.style.font.size = Pt(12)

    doc.save(filename)
    return filename
